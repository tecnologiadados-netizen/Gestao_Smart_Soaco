#!/usr/bin/env python3
"""Preenche o modelo Word do RNC e gera PDF."""

from __future__ import annotations

import argparse
import json
import re
import sys
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document

SCRIPT_DIR = Path(__file__).resolve().parent
SGQ_DIR = SCRIPT_DIR.parent
if str(SGQ_DIR) not in sys.path:
    sys.path.insert(0, str(SGQ_DIR))
from docx_cells import definir_celula

TEMPLATES_DIR = SCRIPT_DIR / "templates"

STATUS_LABELS = {
    "aberto": "Aberto",
    "em_tratamento": "Em andamento",
    "encerrado": "Fechado",
}


def normalizar_texto(texto: str) -> str:
    if not texto:
        return texto
    correcoes = (
        ("BALC\u00c3O", "BALCÃO"),
        ("BALC\u00c3o", "Balcão"),
        ("SERVI\u00c3O", "SERVIÇO"),
        ("Servi\u00c3o", "Serviço"),
        ("servi\u00c3o", "serviço"),
        ("N\u00c3O", "NÃO"),
        ("n\u00c3o", "não"),
    )
    for errado, certo in correcoes:
        texto = texto.replace(errado, certo)
    return texto


def valor_campo(valor: Any) -> str:
    if valor is None:
        return ""
    return normalizar_texto(str(valor).strip())


def formatar_data(valor: Any) -> str:
    texto = valor_campo(valor)
    if not texto:
        return ""
    try:
        if texto.endswith("Z"):
            texto = texto[:-1] + "+00:00"
        data = datetime.fromisoformat(texto)
        return data.strftime("%d/%m/%Y")
    except ValueError:
        return texto


def descricao_ocorrencia(rnc: dict[str, Any]) -> str:
    partes: list[str] = []
    tipo_produto = valor_campo(rnc.get("tipoProduto"))
    descricao = valor_campo(rnc.get("descricaoOcorrencia"))

    if tipo_produto:
        partes.append(f"Tipo de produto: {tipo_produto}")
    if descricao:
        partes.append(descricao)
    return "\n".join(partes)


def codigo_produto(rnc: dict[str, Any]) -> str:
    codigo = valor_campo(rnc.get("codigoProduto"))
    if codigo:
        return codigo
    produto = valor_campo(rnc.get("produto"))
    if not produto:
        return ""
    match = re.match(r"^([A-Z]{2,4}\s+[\dA-Za-z./]+)\s*-", produto, re.IGNORECASE)
    return match.group(1).strip().upper() if match else ""


def descricao_produto(rnc: dict[str, Any]) -> str:
    produto = valor_campo(rnc.get("produto"))
    if not produto:
        return ""
    codigo = codigo_produto(rnc)
    if codigo and produto.upper().startswith(codigo.upper()):
        resto = produto[len(codigo) :].strip()
        if resto.startswith("-"):
            resto = resto[1:].strip()
        return resto or produto
    return produto


STATUS_ACAO_LABELS = {
    "cancelada": "Cancelada",
    "concluida": "Concluída",
    "reprogramada": "Reprogramada",
}


def acoes_apartadas(rnc: dict[str, Any]) -> list[dict[str, Any]]:
    acoes = rnc.get("acoesApartadas")
    if isinstance(acoes, list) and acoes:
        return [item for item in acoes if isinstance(item, dict)]

    legado: list[dict[str, Any]] = []
    pares = (
        ("acaoCorretiva2", "responsavelAcao2", "prazoAcao2"),
        ("acaoCorretiva3", "responsavelAcao3", "prazoAcao3"),
    )
    for acao_key, resp_key, prazo_key in pares:
        acao = valor_campo(rnc.get(acao_key))
        responsavel = valor_campo(rnc.get(resp_key))
        prazo = valor_campo(rnc.get(prazo_key))
        if acao or responsavel or prazo:
            legado.append(
                {
                    "acao": acao,
                    "responsavel": responsavel,
                    "prazoExecucao": prazo,
                    "status": "",
                }
            )
    return legado


def texto_acao_com_status(acao: dict[str, Any]) -> str:
    texto = valor_campo(acao.get("acao"))
    status = valor_campo(acao.get("status"))
    status_label = STATUS_ACAO_LABELS.get(status, status)
    if status_label and texto:
        return f"{texto} (Status: {status_label})"
    if status_label:
        return f"Status: {status_label}"
    return texto


def porques_causa(rnc: dict[str, Any]) -> list[str]:
    porques = rnc.get("porques")
    if isinstance(porques, list):
        valores = [
            valor_campo(item) if isinstance(item, str) else valor_campo(item)
            for item in porques
        ]
        if any(valores):
            return (valores[:5] + [""] * 5)[:5]

    causa = valor_campo(rnc.get("causa"))
    if not causa:
        return ["", "", "", "", ""]
    linhas = [linha.strip() for linha in causa.splitlines() if linha.strip()]
    if len(linhas) >= 2:
        return (linhas[:5] + [""] * 5)[:5]
    return ["", "", "", "", ""]


def montar_campos(payload: dict[str, Any]) -> dict[str, str]:
    registro = payload.get("registro") or {}
    rnc = registro.get("rnc") or {}
    codigo = valor_campo(
        registro.get("codigoDocumento")
        or rnc.get("codigoDocumento")
        or registro.get("numero")
    )
    status = STATUS_LABELS.get(
        valor_campo(registro.get("status")), valor_campo(registro.get("status"))
    )
    porque = porques_causa(rnc)
    acoes = acoes_apartadas(rnc)
    acao_2 = acoes[0] if len(acoes) > 0 else {}
    acao_3 = acoes[1] if len(acoes) > 1 else {}

    return {
        "numero_rnc": codigo,
        "data_registro": formatar_data(rnc.get("dataOcorrencia")),
        "data_fechamento": formatar_data(rnc.get("dataFechamento")),
        "status": status,
        "codigo_produto": codigo_produto(rnc),
        "descricao_produto": descricao_produto(rnc),
        "tipo_acao": valor_campo(rnc.get("tipoAcao")),
        "tipo_ocorrencia": valor_campo(rnc.get("tipoOcorrencia")),
        "quantidade": valor_campo(rnc.get("quantidade")),
        "setor_ocorrencia": valor_campo(rnc.get("setorOcorrencia")),
        "setor_deteccao": valor_campo(rnc.get("setorDeteccao")),
        "lote_serie": valor_campo(rnc.get("loteSerie")),
        "grupo_produto": valor_campo(rnc.get("grupoProduto")),
        "op_numero": valor_campo(rnc.get("numeroOrdemProducao")),
        "nota_fiscal": valor_campo(rnc.get("notaFiscal")),
        "descricao_ocorrencia": descricao_ocorrencia(rnc),
        "preenchido_por": valor_campo(rnc.get("responsavel") or rnc.get("usuarioCriacao")),
        "data_ocorrencia": formatar_data(rnc.get("dataOcorrencia")),
        "acao_imediata": valor_campo(rnc.get("acaoImediata")),
        "descricao_acao_imediata": valor_campo(rnc.get("descricaoAcaoImediata")),
        "responsavel_acao_imediata": valor_campo(rnc.get("responsavelAcaoImediata")),
        "prazo_execucao": formatar_data(rnc.get("prazoExecucao")),
        "abertura_analise_causa": valor_campo(rnc.get("analiseProblema")),
        "porque_1": porque[0],
        "porque_2": porque[1],
        "porque_3": porque[2],
        "porque_4": porque[3],
        "porque_5": porque[4],
        "causa_raiz": valor_campo(rnc.get("causa")),
        "acao_1": valor_campo(rnc.get("resolucaoNaoConformidade")),
        "acao_1_responsavel": valor_campo(rnc.get("responsavelAcaoImediata")),
        "acao_1_prazo": formatar_data(rnc.get("prazoExecucao")),
        "acao_2": texto_acao_com_status(acao_2),
        "acao_2_responsavel": valor_campo(acao_2.get("responsavel")),
        "acao_2_prazo": formatar_data(acao_2.get("prazoExecucao")),
        "acao_3": texto_acao_com_status(acao_3),
        "acao_3_responsavel": valor_campo(acao_3.get("responsavel")),
        "acao_3_prazo": formatar_data(acao_3.get("prazoExecucao")),
        "analise_eficaz": valor_campo(rnc.get("analiseEficaz")),
    }


def preencher_rnc(table, campos: dict[str, str]) -> None:
    definir_celula(table, 1, 1, campos["numero_rnc"])
    definir_celula(table, 1, 7, campos["data_registro"])
    definir_celula(table, 1, 13, campos["data_fechamento"])
    definir_celula(table, 1, 16, campos["status"])
    definir_celula(table, 3, 1, campos["codigo_produto"])
    definir_celula(table, 3, 12, campos["descricao_produto"])
    definir_celula(table, 4, 1, campos["tipo_acao"])
    definir_celula(table, 4, 8, campos["tipo_ocorrencia"])
    definir_celula(table, 4, 13, campos["quantidade"])
    definir_celula(table, 5, 1, campos["setor_ocorrencia"])
    definir_celula(table, 5, 8, campos["setor_deteccao"])
    definir_celula(table, 5, 13, campos["lote_serie"])
    definir_celula(table, 6, 1, campos["grupo_produto"])
    definir_celula(table, 6, 8, campos["op_numero"])
    definir_celula(table, 6, 13, campos["nota_fiscal"])
    definir_celula(table, 8, 0, campos["descricao_ocorrencia"])
    definir_celula(table, 9, 3, campos["preenchido_por"])
    definir_celula(table, 9, 14, campos["data_ocorrencia"])
    definir_celula(table, 11, 3, campos["acao_imediata"])
    definir_celula(table, 13, 0, campos["descricao_acao_imediata"])
    definir_celula(table, 14, 3, campos["responsavel_acao_imediata"])
    definir_celula(table, 14, 14, campos["prazo_execucao"])
    definir_celula(table, 15, 3, campos["abertura_analise_causa"])
    definir_celula(table, 17, 2, campos["porque_1"])
    definir_celula(table, 18, 2, campos["porque_2"])
    definir_celula(table, 19, 2, campos["porque_3"])
    definir_celula(table, 20, 2, campos["porque_4"])
    definir_celula(table, 21, 2, campos["porque_5"])
    definir_celula(table, 22, 2, campos["causa_raiz"])
    definir_celula(table, 23, 2, campos["acao_1"])
    definir_celula(table, 23, 11, campos["acao_1_responsavel"])
    definir_celula(table, 23, 16, campos["acao_1_prazo"])
    definir_celula(table, 24, 2, campos["acao_2"])
    definir_celula(table, 24, 11, campos["acao_2_responsavel"])
    definir_celula(table, 24, 16, campos["acao_2_prazo"])
    definir_celula(table, 25, 2, campos["acao_3"])
    definir_celula(table, 25, 11, campos["acao_3_responsavel"])
    definir_celula(table, 25, 16, campos["acao_3_prazo"])
    definir_celula(table, 26, 2, campos["analise_eficaz"])


def preencher_documento(campos: dict[str, str]) -> Document:
    template_path = TEMPLATES_DIR / "rnc.docx"
    if not template_path.exists():
        raise FileNotFoundError(f"Modelo não encontrado: {template_path}")

    doc = Document(template_path)
    if not doc.tables:
        raise ValueError("O modelo não contém tabelas.")

    preencher_rnc(doc.tables[0], campos)
    return doc


def converter_para_pdf(docx_path: Path, pdf_path: Path) -> None:
    from docx_to_pdf import converter_docx_para_pdf

    converter_docx_para_pdf(str(docx_path), str(pdf_path))


def gerar_pdf(payload: dict[str, Any], output_path: Path) -> None:
    campos = montar_campos(payload)
    doc = preencher_documento(campos)

    with tempfile.TemporaryDirectory() as tmp_dir:
        docx_path = Path(tmp_dir) / "rnc_preenchido.docx"
        doc.save(docx_path)
        converter_para_pdf(docx_path, output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Gera PDF do RNC a partir do modelo Word.")
    parser.add_argument(
        "--input",
        "-i",
        help="Arquivo JSON com os dados do registro. Se omitido, lê stdin.",
    )
    parser.add_argument(
        "--output",
        "-o",
        required=True,
        help="Caminho do arquivo PDF de saída.",
    )
    args = parser.parse_args()

    try:
        if args.input:
            raw = Path(args.input).read_text(encoding="utf-8")
        else:
            raw = sys.stdin.buffer.read().decode("utf-8")
        payload = json.loads(raw)
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        gerar_pdf(payload, output_path)
    except Exception as exc:  # noqa: BLE001
        print(str(exc), file=sys.stderr)
        return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
