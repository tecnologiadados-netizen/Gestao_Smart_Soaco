"""Utilitários compartilhados para preenchimento de células nos modelos Word."""

from __future__ import annotations

from docx.shared import Pt

FONT_NAME = "Arial"
FONT_SIZE = Pt(8)


def formatar_texto_celula(texto: str) -> str:
    """Adiciona margem visual à esquerda para não colar na grade."""
    valor = texto.strip()
    if not valor:
        return ""
    return f" {valor}"


def definir_celula(table, linha: int, coluna: int, texto: str) -> None:
    """Preenche apenas o valor, preservando mesclagens e formatação do modelo."""
    texto_formatado = formatar_texto_celula(texto)
    if not texto_formatado:
        return

    cell = table.cell(linha, coluna)
    if not cell.paragraphs:
        cell.add_paragraph()
    paragraph = cell.paragraphs[0]

    if paragraph.runs:
        paragraph.runs[0].text = texto_formatado
        for run in paragraph.runs[1:]:
            run.text = ""
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run(texto_formatado)

    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE

    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
