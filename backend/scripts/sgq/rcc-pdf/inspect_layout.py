from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

SCRIPT_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = SCRIPT_DIR / "templates"

FONT_NAME = "Arial"
FONT_SIZE_PT = 8


def inspect_doc(path: Path) -> None:
    doc = Document(path)
    print(f"=== {path.name} ===")
    print("sections:", len(doc.sections))
    for i, section in enumerate(doc.sections):
        print(
            f"  S{i}: page {section.page_width.mm:.1f}x{section.page_height.mm:.1f}mm",
            f"margin T{section.top_margin.mm} B{section.bottom_margin.mm}",
            f"L{section.left_margin.mm} R{section.right_margin.mm}",
        )
        print(f"    diff_first={section.different_first_page_header_footer}")
    print("paragraphs:", len(doc.paragraphs))
    for pi, p in enumerate(doc.paragraphs[:5]):
        print(f"  P{pi}: '{p.text[:60]}'")
    print("tables:", len(doc.tables))
    if doc.tables:
        table = doc.tables[0]
        total_height = 0
        for ri, row in enumerate(table.rows):
            h = row.height
            total_height += h.twips if h else 0
            if ri in {0, 1, 13, 14, 17, 19, 20, 21}:
                print(f"  R{ri} height={h}")
        print(f"  approx total row height twips: {total_height}")
    for rel in doc.part.rels.values():
        if "header" in rel.reltype or "footer" in rel.reltype:
            print("  rel:", rel.reltype, rel.target_ref)


for name in ["cliente.docx", "empresa.docx"]:
    inspect_doc(TEMPLATES_DIR / name)
    print()
