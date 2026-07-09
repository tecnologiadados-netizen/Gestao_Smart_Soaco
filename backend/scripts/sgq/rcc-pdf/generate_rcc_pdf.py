#!/usr/bin/env python3
"""Preenche o modelo Word do RCC e gera PDF (cliente ou empresa)."""

from __future__ import annotations

import argparse
import json
import sys
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

SCRIPT_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = SCRIPT_DIR / "templates"

FONT_NAME = "Arial"
FONT_SIZE = Pt(8)
VENDEDOR_PADRAO = "SO AÇO INDUSTRIAL LTDA - SO MOVEIS LTDA"

STATUS_LABELS = {
    "aberto": "Aberto",
    "em_tratamento": "Em andamento",
    "encerrado": "Fechado",
}


def normalizar_texto(texto: str) -> str:
    """Corrige caracteres quebrados vindos do histórico ERP/Nomus."""
    if not texto:
        return texto
    correcoes = (
        ("BALC\u00c3O", "BALCÃO"),
        ("BALC\u00c3o", "Balcão"),
        ("SERVI\u00c3O", "SERVIÇO"),
        ("Servi\u00c3o", "Serviço"),
        ("servi\u00c3o", "serviço"),
        ("N\u00c3O", "NÃO"),
        ("n\u00c3o", "não"),
    )
    for errado, certo in correcoes:
        texto = texto.replace(errado, certo)
    return texto


def valor_campo(valor: Any) -> str:
    if valor is None:
        return ""
    return normalizar_texto(str(valor).strip())


def formatar_data(valor: Any) -> str:
    texto = valor_campo(valor)
    if not texto:
        return ""
    try:
        if texto.endswith("Z"):
            texto = texto[:-1] + "+00:00"
        data = datetime.fromisoformat(texto)
        return data.strftime("%d/%m/%Y")
    except ValueError:
        return texto


def formatar_hora(valor: Any) -> str:
    texto = valor_campo(valor)
    if not texto:
        return ""
    if len(texto) >= 5 and ":" in texto:
        return texto[:5]
    return texto


def nome_revendedor(rcc: dict[str, Any]) -> str:
    if rcc.get("clienteDoRevendedor"):
        return valor_campo(rcc.get("nomeRevendedor"))
    return valor_campo(rcc.get("vendedor")) or VENDEDOR_PADRAO


def tipo_reclamacao(rcc: dict[str, Any]) -> str:
    return valor_campo(rcc.get("reclamacao1") or rcc.get("reclamacao2"))


def servicos_realizados(rcc: dict[str, Any]) -> str:
    partes = [
        valor_campo(rcc.get("servicoRealizado")),
        valor_campo(rcc.get("servicoRealizado1")),
        valor_campo(rcc.get("servicoRealizado2")),
    ]
    return " · ".join(p for p in partes if p)


def montar_campos(payload: dict[str, Any]) -> dict[str, str]:
    registro = payload.get("registro") or {}
    rcc = registro.get("rcc") or {}
    codigo = valor_campo(
        registro.get("codigoDocumento")
        or rcc.get("codigoDocumento")
        or registro.get("numero")
    )
    status = STATUS_LABELS.get(
        valor_campo(registro.get("status")), valor_campo(registro.get("status"))
    )

    cidade = valor_campo(rcc.get("cidade"))
    estado = valor_campo(rcc.get("estado"))
    if cidade and estado and estado not in cidade:
        cidade = f"{cidade}-{estado}"

    return {
        "numero_reclamacao": codigo,
        "data_registro": formatar_data(rcc.get("dataRegistroReclamacao")),
        "data_fechamento": formatar_data(rcc.get("dataFechamento")),
        "status": status,
        "nome_consumidor": valor_campo(rcc.get("nomeClienteConsumidor")),
        "nome_revendedor": nome_revendedor(rcc),
        "contato": valor_campo(rcc.get("contato")),
        "cidade": cidade,
        "telefone": valor_campo(rcc.get("telefone")),
        "bairro": valor_campo(rcc.get("bairro")),
        "endereco": valor_campo(rcc.get("endereco")),
        "ponto_referencia": valor_campo(rcc.get("pontoReferencia")),
        "produto": valor_campo(rcc.get("produto")),
        "serie_lote": valor_campo(rcc.get("numeroSerieLoteProduto")),
        "data_nf": formatar_data(rcc.get("dataEmissaoNf")),
        "nota_fiscal": valor_campo(rcc.get("numeroNf")),
        "quantidade": valor_campo(rcc.get("quantidade")),
        "pedido": valor_campo(rcc.get("numeroPedidoInternoExterno")),
        "tipo_reclamacao": tipo_reclamacao(rcc),
        "descricao_reclamacao": valor_campo(rcc.get("descricaoReclamacao")),
        "reclamacao_aceita": valor_campo(rcc.get("reclamacaoAceita")),
        "dentro_garantia": valor_campo(rcc.get("produtoDentroGarantia")),
        "abrir_os": valor_campo(rcc.get("abrirOrdemServico")),
        "comentario": valor_campo(rcc.get("comentario")),
        "responsavel_analise": valor_campo(rcc.get("usuarioCriacao")),
        "ordem_producao": valor_campo(rcc.get("numeroOrdemProducao")),
        "data_assistencia": formatar_data(rcc.get("dataAssistencia")),
        "funcionario": valor_campo(rcc.get("funcionarioSolicitado")),
        "hora_saida_empresa": formatar_hora(rcc.get("horaSaidaEmpresa")),
        "serie_compressor": valor_campo(rcc.get("numeroSerieCompressor")),
        "hora_chegada_empresa": formatar_hora(rcc.get("horaChegadaEmpresa")),
        "servico_realizado": servicos_realizados(rcc),
        "problema_solucionado": valor_campo(rcc.get("problemaSolucionado")),
        "data_conclusao": formatar_data(rcc.get("dataFechamento")),
    }


def definir_celula(table, linha: int, coluna: int, texto: str) -> None:
    """Preenche apenas o valor, preservando mesclagens e formatação do modelo."""
    if not texto:
        return

    cell = table.cell(linha, coluna)
    if not cell.paragraphs:
        cell.add_paragraph()
    paragraph = cell.paragraphs[0]

    if paragraph.runs:
        paragraph.runs[0].text = texto
        for run in paragraph.runs[1:]:
            run.text = ""
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run(texto)

    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE

    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def preencher_cliente(table, campos: dict[str, str]) -> None:
    definir_celula(table, 1, 1, campos["numero_reclamacao"])
    definir_celula(table, 1, 6, campos["data_registro"])
    definir_celula(table, 1, 11, campos["data_fechamento"])
    definir_celula(table, 1, 15, campos["status"])
    definir_celula(table, 3, 1, campos["nome_consumidor"])
    definir_celula(table, 3, 13, campos["nome_revendedor"])
    definir_celula(table, 4, 1, campos["contato"])
    definir_celula(table, 4, 11, campos["cidade"])
    definir_celula(table, 5, 1, campos["telefone"])
    definir_celula(table, 5, 11, campos["bairro"])
    definir_celula(table, 6, 1, campos["endereco"])
    definir_celula(table, 6, 11, campos["ponto_referencia"])
    definir_celula(table, 8, 3, campos["produto"])
    definir_celula(table, 8, 12, campos["serie_lote"])
    definir_celula(table, 9, 3, campos["data_nf"])
    definir_celula(table, 9, 12, campos["nota_fiscal"])
    definir_celula(table, 10, 3, campos["quantidade"])
    definir_celula(table, 10, 12, campos["pedido"])
    definir_celula(table, 12, 7, campos["tipo_reclamacao"])
    definir_celula(table, 14, 0, campos["descricao_reclamacao"])
    definir_celula(table, 16, 2, campos["reclamacao_aceita"])
    definir_celula(table, 16, 7, campos["dentro_garantia"])
    definir_celula(table, 16, 14, campos["abrir_os"])
    definir_celula(table, 17, 2, campos["comentario"])
    definir_celula(table, 18, 2, campos["responsavel_analise"])


def preencher_empresa(table, campos: dict[str, str]) -> None:
    definir_celula(table, 1, 1, campos["numero_reclamacao"])
    definir_celula(table, 1, 5, campos["data_registro"])
    definir_celula(table, 1, 10, campos["data_fechamento"])
    definir_celula(table, 1, 15, campos["status"])
    definir_celula(table, 3, 1, campos["nome_consumidor"])
    definir_celula(table, 3, 11, campos["nome_revendedor"])
    definir_celula(table, 4, 1, campos["contato"])
    definir_celula(table, 4, 10, campos["cidade"])
    definir_celula(table, 5, 1, campos["telefone"])
    definir_celula(table, 5, 10, campos["bairro"])
    definir_celula(table, 6, 1, campos["endereco"])
    definir_celula(table, 6, 10, campos["ponto_referencia"])
    definir_celula(table, 8, 3, campos["produto"])
    definir_celula(table, 8, 12, campos["serie_lote"])
    definir_celula(table, 9, 3, campos["data_nf"])
    definir_celula(table, 9, 12, campos["nota_fiscal"])
    definir_celula(table, 10, 3, campos["quantidade"])
    definir_celula(table, 10, 12, campos["pedido"])
    definir_celula(table, 12, 3, campos["tipo_reclamacao"])
    definir_celula(table, 14, 0, campos["descricao_reclamacao"])
    definir_celula(table, 17, 4, campos["ordem_producao"])
    definir_celula(table, 17, 13, campos["data_assistencia"])
    definir_celula(table, 18, 4, campos["funcionario"])
    definir_celula(table, 18, 13, campos["hora_saida_empresa"])
    definir_celula(table, 19, 4, campos["serie_compressor"])
    definir_celula(table, 19, 13, campos["hora_chegada_empresa"])
    definir_celula(table, 21, 4, campos["servico_realizado"])
    definir_celula(table, 22, 4, campos["problema_solucionado"])
    definir_celula(table, 22, 14, campos["data_conclusao"])


def preencher_documento(versao: str, campos: dict[str, str]) -> Document:
    template_name = "cliente.docx" if versao == "cliente" else "empresa.docx"
    template_path = TEMPLATES_DIR / template_name
    if not template_path.exists():
        raise FileNotFoundError(f"Modelo não encontrado: {template_path}")

    doc = Document(template_path)
    if not doc.tables:
        raise ValueError("O modelo não contém tabelas.")

    table = doc.tables[0]
    if versao == "cliente":
        preencher_cliente(table, campos)
    else:
        preencher_empresa(table, campos)
    return doc


def converter_para_pdf(docx_path: Path, pdf_path: Path) -> None:
    try:
        from docx2pdf import convert
    except ImportError as exc:
        raise RuntimeError(
            "Pacote docx2pdf não instalado. Execute: pip install -r scripts/rcc-pdf/requirements.txt"
        ) from exc

    convert(str(docx_path), str(pdf_path))


def gerar_pdf(payload: dict[str, Any], output_path: Path) -> None:
    versao = valor_campo(payload.get("versao")).lower()
    if versao not in {"cliente", "empresa"}:
        raise ValueError("Versão inválida. Use 'cliente' ou 'empresa'.")

    campos = montar_campos(payload)
    doc = preencher_documento(versao, campos)

    with tempfile.TemporaryDirectory() as tmp_dir:
        docx_path = Path(tmp_dir) / "rcc_preenchido.docx"
        doc.save(docx_path)
        converter_para_pdf(docx_path, output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Gera PDF do RCC a partir do modelo Word.")
    parser.add_argument(
        "--input",
        "-i",
        help="Arquivo JSON com os dados do registro. Se omitido, lê stdin.",
    )
    parser.add_argument(
        "--output",
        "-o",
        required=True,
        help="Caminho do arquivo PDF de saída.",
    )
    args = parser.parse_args()

    try:
        if args.input:
            raw = Path(args.input).read_text(encoding="utf-8")
        else:
            raw = sys.stdin.buffer.read().decode("utf-8")
        payload = json.loads(raw)
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        gerar_pdf(payload, output_path)
    except Exception as exc:  # noqa: BLE001
        print(str(exc), file=sys.stderr)
        return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
