"""
Preenche o template Word e exporta para PDF via Microsoft Word.
O layout da página não é alterado — exportação idêntica a 'Salvar como PDF' no Word.
"""
from __future__ import annotations

import os
import re
import sys
import tempfile
from copy import deepcopy
from datetime import date, datetime
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

BACKEND_ROOT = Path(__file__).resolve().parent.parent.parent
SGQ_DIR = BACKEND_ROOT / "scripts" / "sgq"
if str(SGQ_DIR) not in sys.path:
    sys.path.insert(0, str(SGQ_DIR))
from docx_to_pdf import converter_docx_para_pdf

TEMPLATE_PATH = BACKEND_ROOT / "assets" / "pre-compra" / "formulario_cotacao.docx"

FONT_NAME = "Arial"
FONT_SIZE = Pt(8)
HEADER_PAD_LEFT = Pt(6)

ITEM_ROW_IDX = 8
SOLICITACAO_VALUE_COL = 3


def _fmt_date(value) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    text = str(value).strip()
    if "T" in text:
        try:
            return datetime.fromisoformat(text).strftime("%d/%m/%Y")
        except ValueError:
            pass
    if len(text) >= 10:
        try:
            return datetime.strptime(text[:10], "%Y-%m-%d").strftime("%d/%m/%Y")
        except ValueError:
            pass
    return text


def _fmt_date_short(value) -> str:
    if value is None:
        return ""
    if isinstance(value, (datetime, date)):
        return value.strftime("%d/%m/%y")
    text = str(value).strip()
    if "T" in text:
        try:
            return datetime.fromisoformat(text).strftime("%d/%m/%y")
        except ValueError:
            pass
    if len(text) >= 10:
        try:
            return datetime.strptime(text[:10], "%Y-%m-%d").strftime("%d/%m/%y")
        except ValueError:
            pass
    return text


def _fmt_money(value) -> str:
    if value is None:
        return "R$ 0,00"
    amount = Decimal(str(value))
    formatted = f"{amount:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"R$ {formatted}"


def _fmt_qty(value) -> str:
    if value is None or value == "":
        return ""
    amount = Decimal(str(value))
    return f"{amount:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").lower().strip())


def _apply_font_cell(cell, *, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE
            run.font.bold = bold


def _set_cell(
    row,
    index: int,
    text,
    *,
    center: bool = False,
    pad_left: bool = False,
    bold: bool = False,
) -> None:
    cell = row.cells[index]
    value = str(text or "")

    if cell.paragraphs:
        first = cell.paragraphs[0]
        if first.runs:
            first.runs[0].text = value
            for run in first.runs[1:]:
                run.text = ""
        else:
            run = first.add_run(value)
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE
        for extra in cell.paragraphs[1:]:
            for run in extra.runs:
                run.text = ""
    else:
        cell.text = value

    for paragraph in cell.paragraphs:
        if center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if pad_left:
            paragraph.paragraph_format.left_indent = HEADER_PAD_LEFT

    _apply_font_cell(cell, bold=bold)


def _duplicate_row_after(table, row_index: int) -> None:
    template_row = table.rows[row_index]
    new_tr = deepcopy(template_row._tr)
    template_row._tr.addnext(new_tr)


def _find_row(table, fragment: str) -> int | None:
    needle = _normalize_text(fragment)
    for i, row in enumerate(table.rows):
        if any(needle in _normalize_text(cell.text) for cell in row.cells):
            return i
    return None


def _fill_item_row(row, item: dict) -> None:
    values = {
        0: item.get("codigo_produto"),
        1: item.get("codigo_fornecedor"),
        3: item.get("descricao_produto"),
        5: _fmt_qty(item.get("qtde")),
        7: item.get("unidade"),
        9: _fmt_money(item.get("preco_unitario")),
        12: _fmt_money(item.get("valor_total")),
    }
    for col, value in values.items():
        _set_cell(row, col, value, center=True)


def _fmt_coleta(data: dict) -> str:
    """Nº da coleta: aceita string única, número ou lista (`numeros_coleta`)."""
    for key in ("coleta", "numero_coleta"):
        value = data.get(key)
        if value not in (None, ""):
            return str(value)
    numeros = data.get("numeros_coleta")
    if isinstance(numeros, (list, tuple)):
        return ", ".join(str(n) for n in numeros if n not in (None, ""))
    if numeros not in (None, ""):
        return str(numeros)
    return ""


def _cell_text(row, index: int) -> str:
    if index < 0 or index >= len(row.cells):
        return ""
    return _normalize_text(row.cells[index].text)


def _find_header_value_cell(table, *label_fragments: str) -> tuple[int, int] | None:
    """Localiza (linha, coluna_valor) pelo texto do rótulo; valor = primeira célula distinta à direita."""
    needles = [_normalize_text(f) for f in label_fragments if f]
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            label = _normalize_text(cell.text)
            if not label or not any(n in label for n in needles):
                continue
            for k in range(j + 1, len(row.cells)):
                if _normalize_text(row.cells[k].text) != label:
                    return i, k
            return i, min(j + 1, len(row.cells) - 1)
    return None


def _fill_header(table, data: dict) -> None:
    # Preferência: achar pelo rótulo do template (aceita nome antigo ou novo).
    # Fallback: índices fixos do layout atual do formulario_cotacao.docx.
    resolved: list[tuple[int, int, object]] = []
    specs: list[tuple[tuple[str, ...], int, int, object]] = [
        (
            ("nº da pré-compra", "n° da pré-compra", "nº da pre-compra", "n° da pre-compra", "nº da cotação", "n° da cotação", "nº da cotacao", "n° da cotacao"),
            1,
            2,
            data.get("cotacao"),
        ),
        (("nº da coleta", "n° da coleta", "nº da coleta", "n° da coleta"), 1, 6, _fmt_coleta(data)),
        (("data de emissão", "data de emissao"), 1, 11, _fmt_date(data.get("data_emissao"))),
        (("nome do comprador",), 2, 2, data.get("comprador")),
        (("telefone comprador",), 2, 10, data.get("telefone")),
        (("fornecedor:", "fornecedor"), 3, 2, data.get("fornecedor")),
        (("cnpj",), 3, 8, data.get("cnpj")),
        (("contato:", "contato"), 4, 2, data.get("contato")),
        (("telefone:", "telefone"), 4, 8, data.get("telefone_fornecedor")),
    ]
    used_rows_cols: set[tuple[int, int]] = set()
    for fragments, fb_row, fb_col, value in specs:
        found = _find_header_value_cell(table, *fragments)
        if found is not None and found not in used_rows_cols:
            row_idx, col_idx = found
        else:
            row_idx, col_idx = fb_row, fb_col
        if (row_idx, col_idx) in used_rows_cols:
            row_idx, col_idx = fb_row, fb_col
        used_rows_cols.add((row_idx, col_idx))
        resolved.append((row_idx, col_idx, value))

    for row_idx, col_idx, value in resolved:
        if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
            _set_cell(table.rows[row_idx], col_idx, value, pad_left=True, bold=False)


def _fmt_solicitacao(sol: dict) -> str:
    num = sol.get("id")
    data = _fmt_date_short(sol.get("data_necessidade"))
    if num and data:
        return f"{num} - {data}"
    return str(num or data or "")


def _fill_solicitacoes(table, solicitacoes: list[dict]) -> None:
    sol_idx = _find_row(table, "data de necessidade")
    if sol_idx is None:
        sol_idx = _find_row(table, "solicita")
    if sol_idx is None:
        return

    ordenadas = sorted(solicitacoes, key=lambda s: s.get("id") or 0)

    if not ordenadas:
        _set_cell(table.rows[sol_idx], SOLICITACAO_VALUE_COL, "", pad_left=True)
        return

    if len(ordenadas) == 1:
        texto = _fmt_solicitacao(ordenadas[0])
    else:
        texto = "; ".join(
            f"{i}° {_fmt_solicitacao(sol)}" for i, sol in enumerate(ordenadas, start=1)
        )

    _set_cell(table.rows[sol_idx], SOLICITACAO_VALUE_COL, texto, pad_left=True, bold=False)


def _fill_document(data: dict) -> Document:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template não encontrado: {TEMPLATE_PATH}")

    doc = Document(str(TEMPLATE_PATH))
    table = doc.tables[0]

    _fill_header(table, data)

    items = data.get("itens") or [{}]
    _fill_item_row(table.rows[ITEM_ROW_IDX], items[0])
    current = ITEM_ROW_IDX
    for item in items[1:]:
        _duplicate_row_after(table, current)
        current += 1
        _fill_item_row(table.rows[current], item)

    total_row_idx = _find_row(table, "VALOR TOTAL")
    if total_row_idx is not None:
        _set_cell(table.rows[total_row_idx], 12, _fmt_money(data.get("valor_total_geral")), center=True)

    _fill_solicitacoes(table, data.get("solicitacoes") or [])
    return doc


def generate_pdf(data: dict) -> bytes:
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "cotacao.docx")
        pdf_path = os.path.join(tmp, "cotacao.pdf")

        doc = _fill_document(data)
        doc.save(docx_path)
        converter_docx_para_pdf(docx_path, pdf_path)

        with open(pdf_path, "rb") as f:
            return f.read()
