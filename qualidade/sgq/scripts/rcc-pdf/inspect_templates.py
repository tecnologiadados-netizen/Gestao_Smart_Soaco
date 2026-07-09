from pathlib import Path

from docx import Document

DOWNLOADS = Path(r"c:\Users\Davi\Downloads")
TEMPLATES = [
    DOWNLOADS / "formulario_RCC_V.cliente.docx",
    DOWNLOADS / "formulario_RCC_V.Empresa.docx",
]

for path in TEMPLATES:
    doc = Document(path)
    print(f"=== {path.name} | tables: {len(doc.tables)} ===")
    for ti, table in enumerate(doc.tables):
        print(f"-- Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
        for ri, row in enumerate(table.rows):
            cells = [c.text.replace("\n", " / ").strip() for c in row.cells]
            if any(cells):
                print(f"  R{ri}: {' | '.join(cells)}")
    print()
